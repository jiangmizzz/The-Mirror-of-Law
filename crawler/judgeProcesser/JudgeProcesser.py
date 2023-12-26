from elasticsearch import Elasticsearch
import openpyxl
import re
import docx
import warnings  

warnings.filterwarnings("ignore")
class JudgeProcesser():
    def __init__(self):
        self.rootPath = "./judgeData/"
        self.file = "河南行政.xlsx"
        pass

    def readExcel(self, fileName):
        excel = openpyxl.load_workbook(self.rootPath + fileName)
        sheet = excel.active
        return sheet
    
    def processData(self):
        sheet = self.readExcel(self.file)
        for row in sheet.iter_rows():
            if(row[1].value == None or row[1].value == "标题"):
                continue
            judgeDocument = JudgeDocument(row[1].value, row[2].value, row[3].value, row[4].value,
                           row[5].value, row[6].value, row[7].value, row[8].value, row[9].value)
            order = int(row[0].value)
            if(order == 18):
                print(111)
            content = []
            try:
                document = docx.Document(self.rootPath + self.file.split('.')[0] + '/' +
                                          str(order) + '-' + judgeDocument.title.replace(':', ' ') + '.docx')
                for paragraph in document.paragraphs:
                    if(paragraph.text == '' or paragraph.text == '更多信息请点击查看把手案例'):
                        continue
                    content.append(paragraph.text.replace("\u3000", " "))
            except:
                continue
            judgeDocument.content = '\n'.join(content)
            judgeDocument.index()

class JudgeDocument():
    def __init__(self, link, caseType, cause, procedure, docType, province, court, caseId, date):
        self.title = re.search(title_pattern, link).group(1)
        self.url = re.search(link_pattern, link).group()  
        self.caseType = caseType
        self.cause = cause
        self.procedure = procedure
        self.docType = docType
        self.province = province
        self.court = court
        self.caseId = caseId
        self.date = date
        self.content = None
        self.like = 0
        self.dislike = 0

    def index(self):
        self.title = self.title.replace('null', '')
        document = vars(self)
        try:
            es.index(index=index, document=document, id=re.search(id_pattern, self.url).group(1))
        except:
            return

link_pattern = r'https://www.lawsdata.com/#/documentDetails\?id=[a-zA-Z0-9]+&type=\d+'  
title_pattern = r',"(.*?)"\)' 
id_pattern = r'id=([a-zA-Z0-9]+)'

if __name__ == "__main__":
    es = Elasticsearch("https://ES服务器ip:9200",http_auth=('ES用户名', 'ES密码'), timeout=20, verify_certs=False)
    index = 'judgements'
    judgeProcesser = JudgeProcesser()
    judgeProcesser.processData()