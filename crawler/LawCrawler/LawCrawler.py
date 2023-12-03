import requests
import re
import os
import json
import docx
import win32com.client as wc
from bs4 import BeautifulSoup
from elasticsearch import Elasticsearch # 8.11.0
import warnings  

from Cracker import Cracker, OCR
  
warnings.filterwarnings("ignore")

class LawCrawler():
    def __init__(self):
        self.session = requests.Session()
        self.rootApi = "https://flk.npc.gov.cn/api"
        self.detailUrl = "https://flk.npc.gov.cn/api/detail"
        self.rootUrl = "https://flk.npc.gov.cn"
        self.downloadDocUrl = "https://wb.flk.npc.gov.cn"
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36',
            }
        self.typeDic = {"法律法规": "flfg", "行政法规": "xzfg", "监察法规": "jcfg", "司法解释": "sfjs", "地方性法规" : "dfxfg"}
        self.fail = 0
        self.success = 0
    
    def getRootUrl(self, type, page):
        return self.rootApi + f"/?type={self.typeDic[type]}&searchType=title%3Bvague&sortTr=f_bbrq_s%3Bdesc&sort=true&page={page}&size=10"

    def crawl(self, type, startPage, pageNum):
        if type == "宪法":
            self.crawlXF()
            return
        
        for i in range(startPage, startPage + pageNum):
            # 获取目录页数据
            res = self.session.get(self.getRootUrl(type, i + 1), headers=self.headers)
            if res.text.startswith('{'):
                jsonData = json.loads(res.text)
            else:
                while True:
                    if len(re.findall(r"x=\'(.*?)\'", res.text)):
                        _x = re.findall(r"x=\'(.*?)\'", res.text)[0]
                        _i = re.findall(r"i=\'(.*?)\'", res.text)[0]
                        _d = re.findall(r"d=\'(.*?)\'", res.text)[0]
                        _g = re.findall(r"g=\'(.*?)\'", res.text)[0]
                        cracker = Cracker(_x, _i, _d, _g)
                        crackerUrl = cracker.getUrl()
                        print("trying to crack...")
                        print("crack url = " + crackerUrl)
                        res = self.session.get(self.rootUrl + crackerUrl, headers=self.headers)
                        if res.status_code == 302:
                            res = self.session.get(self.getRootUrl(type, i + 1), headers=self.headers)
                            jsonData = json.loads(res.text)
                            break
                        elif res.text.startswith('{'):
                            jsonData = json.loads(res.text)
                            break
                        else:   
                            continue
                    else:
                        res = self.session.get('https://flk.npc.gov.cn/waf_text_captcha', headers=self.headers)
                        path = './data/captcha.jpg'
                        with open(path, mode='wb') as f:
                            f.write(res.content)
                        captcha = OCR.OCR(path)
                        print("验证码自动识别：" + captcha)
                        res = self.session.get('https://flk.npc.gov.cn/waf_text_verify.html?captcha=' + captcha, headers=self.headers)
                        res = self.session.get(self.getRootUrl(type, i + 1), headers=self.headers)
                        if res.text.startswith('{'):
                            jsonData = json.loads(res.text)
                            break
                        else:
                            print("验证码识别错误，正在重试……")
                            continue

            listData = jsonData['result']['data']
            for item in listData:
                docId = item['id']
                # 获取指定id的文件的下载链接
                res = self.session.post(self.detailUrl, headers=self.headers,
                                        params={'id': docId})

                if res.text.startswith('{'):
                    jsonData = json.loads(res.text)
                else:
                    while True:
                        if len(re.findall(r"x=\'(.*?)\'", res.text)):
                            _x = re.findall(r"x=\'(.*?)\'", res.text)[0]
                            _i = re.findall(r"i=\'(.*?)\'", res.text)[0]
                            _d = re.findall(r"d=\'(.*?)\'", res.text)[0]
                            _g = re.findall(r"g=\'(.*?)\'", res.text)[0]
                            cracker = Cracker(_x, _i, _d, _g)
                            crackerUrl = cracker.getUrl()
                            print("trying to crack...")
                            print("crack url = " + crackerUrl)
                            res = self.session.get(self.rootUrl + crackerUrl, headers=self.headers)
                            if res.status_code == 302:
                                res = self.session.post(self.detailUrl, headers=self.headers,
                                        params={'id': docId})
                                jsonData = json.loads(res.text)
                                break
                            elif res.text.startswith('{'):
                                jsonData = json.loads(res.text)
                                break
                            else:   
                                continue
                        else:
                            res = self.session.get('https://flk.npc.gov.cn/waf_text_captcha', headers=self.headers)
                            path = './data/captcha.jpg'
                            with open(path, mode='wb') as f:
                                f.write(res.content)
                            captcha = OCR.OCR(path)
                            print("验证码自动识别：" + captcha)
                            res = self.session.get('https://flk.npc.gov.cn/waf_text_verify.html?captcha=' + captcha, headers=self.headers)
                            res = self.session.post(self.detailUrl, headers=self.headers,
                                        params={'id': docId})
                            if res.text.startswith('{'):
                                jsonData = json.loads(res.text)
                                break
                            else:
                                print("验证码识别错误，正在重试……")
                                continue

                docUrl = None
                # 获取docx格式文件的url
                for doc in jsonData['result']['body']:
                    if doc['type'] == 'WORD': 
                       docUrl = self.downloadDocUrl + doc['path']
                       break
                
                # 有的文档没有WORD，懒得处理了直接continue
                if docUrl is None:
                    self.fail = self.fail + 1
                    print(str(self.fail) + ": fail to fetch " + jsonData['result']['title'])
                    continue

                self.success = self.success + 1
                print(str(self.success) + ": " + docUrl)
                # 解析文档数据
                title = jsonData['result']['title']
                level = jsonData['result']['level']
                office = jsonData['result']['office']
                publish = jsonData['result']['publish'].split(' ')[0]
                expiry = jsonData['result']['expiry'].split(' ')[0]
                document = LawDocument(title=title, level=level, office=office,
                                        publish=publish, expiry=expiry, url=docUrl, type=type)
                res = self.session.get(docUrl, headers=self.headers)
                path = f'./data/{type}/' + title + '_' + publish.split(' ')[0] + '.' + docUrl.split('.')[-1]
                with open(path, mode='wb') as f:
                    f.write(res.content)

                # 解析docx数据
                if document.parseDoc(path):
                    document.index()
                else:
                    self.fail = self.fail + 1
                    print(str(self.fail) + ": fail to fetch " + jsonData['result']['title'])

    def crawlXF(self): 
        # 宪法总共就7个docx
        for i in range(7):
            # docx的下载链接是写死在js里的
            XFUrl = f"https://flk.npc.gov.cn/xf/js/xf{i + 1}.js"
            infoUrl = f"https://flk.npc.gov.cn/xf/html/xf{i + 1}.html"
            res = self.session.get(XFUrl, headers=self.headers)
            docUrl = re.search(r'downLoadWordFileFileBs="(.*?)"', res.text).group(1)
            print(docUrl)
            res = self.session.get(infoUrl, headers=self.headers)
            soup = BeautifulSoup(res.text, 'html.parser')
            title = soup.find(class_="title").text
            level = soup.find(id="xlwj").text
            office = soup.find(id="zdjg").text
            publish = soup.find(id="gbrq").text.split(' ')[0]
            expiry = soup.find(id="sxrq").text.split(' ')[0]
            document = LawDocument(title=title, level=level, office=office, 
                                   publish=publish, expiry=expiry, url=docUrl, type="宪法")
            res = self.session.get(docUrl, headers=self.headers)
            path = './data/宪法/' + title + '.docx'
            with open(path, mode='wb') as f:
                f.write(res.content)

            # 解析docx数据
            if document.parseDoc(path):
                document.index()
    
    def createDir(self):
        if not os.path.exists("./data"):  
            os.makedirs("./data")  
        dirList = ["./data/宪法", "./data/法律法规", "./data/行政法规", "./data/监察法规", "./data/司法解释", "./data/地方性法规"]
        for dir in dirList:
            if not os.path.exists(dir):  
                os.makedirs(dir)  

class LawDocument():
    def __init__(self, title, level, office, publish, expiry, url, type):
        self.title = title
        self.level = level
        self.office = office
        self.publish = publish
        self.expiry = expiry
        self.url = url
        self.type = type
        self.content = None
        self.like = 0
        self.dislike = 0

    def parseDoc(self, docPath : str):
        try:
            document = docx.Document(docPath)  # 建立Word文件对象
        except:
            if docPath.endswith('doc'):
                try:
                    word = wc.Dispatch("Word.Application")
                    doc = word.Documents.Open(os.path.abspath(docPath))
                    # 12代表转换后为docx文件
                    doc.SaveAs(os.path.abspath(docPath + "x"), 12)
                    doc.Close()
                    word.Quit()
                    document = docx.Document(docPath + "x")
                except:
                    return False
            else:
                return False
        
        content = []
        for paragraph in document.paragraphs:
            content.append(paragraph.text)  # 将每一段Paragraph组成列表
        delLine = 0
        for i in range(min(10, len(content))):
            if content[i].strip() != '' and content[i].strip()[0] == '(':
                break
            else:
                delLine = delLine + 1
        if delLine < min(10, len(content)):
            del content[:delLine]
        else:
            del content[:4]
        self.content = '\n'.join(content)  # 将列表转成字符串并隔行输出
        # print(self.content)
        return True
    
    def index(self):
        document = vars(self)
        if document['publish'] == '' and document['expiry'] != '':
            document['publish'] = document['expiry']
        if document['expiry'] == '' and document['publish'] != '':
            document['expiry'] = document['publish']
        if document['publish'] == '' and document['expiry'] == '':
            document.pop('publish')
            document.pop('expiry')

        try:
            es.index(index='laws', document=document, id=self.url.split('/')[-1].split('.')[0])
        except:
            return
        
if __name__ == "__main__":
    es = Elasticsearch("https://ES服务器ip:9200",http_auth=('ES用户名', 'ES密码'), timeout=20, verify_certs=False)
    index = 'laws'
    # res = es.indices.create(index='laws', ignore=400)
    crawler = LawCrawler()
    crawler.createDir()
    # 参数1可选：宪法，法律法规，行政法规，监察法规，司法解释，地方性法规
    # 参数2为爬取的起始页码，从0开始
    # 参数3为爬取页数，每页10个docx
    crawler.crawl("法律法规", 0, 100)